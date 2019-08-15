import sqlite3
from tkinter import *
import tkinter.ttk as ttk
from docx import Document
import os
from datetime import datetime


class Table(Frame):
    def __init__(self, parent=None, headings=tuple(), rows=tuple()):
        super().__init__(parent)

        self.db_access = PyAccountingDataBase()
        self.orders_generator = WordOrdedGenerator()

        self.table = ttk.Treeview(self, show="headings", selectmode="browse")
        self.table["columns"] = headings
        self.table["displaycolumns"] = headings

        for head in headings:
            self.table.heading(head, text=head, anchor=CENTER)
            self.table.column(head, anchor=CENTER)

        for row in rows:
            self.table.insert('', END, values=tuple(row))

        if headings[1] == 'Ім\'я':
            self.table.bind('<Double-Button-1>', self.sub_employee_menu)
        elif headings[1] == 'Назва':
            self.table.bind('<Double-Button-1>', self.sub_division_menu)

        scrolltable = Scrollbar(self, command=self.table.yview)
        self.table.configure(yscrollcommand=scrolltable.set)
        scrolltable.pack(side=RIGHT, fill=Y)
        self.table.pack(expand=YES, fill=BOTH)

    def sub_employee_menu(self, event):
        self.rt = Toplevel(self.table)
        self.rt.title('Конфігурація працівника')

        selected_item = self.table.selection()[0]
        listed_raw = [str(i) for i in list(self.table.item(selected_item)['values'])]
        name, position, subdivision = ' '.join(listed_raw[1:4]), listed_raw[4], listed_raw[5]
        res = '- ' + name + ';\n- ' + position + ';\n- ' + subdivision + ';'

        self.fr = Frame(self.rt, bd=5)
        self.fr.pack(fill=BOTH, expand=True, side=TOP)

        self.upper_label = Label(self.fr, text='Меню робітника:\n{}'.format(res), font=('arial', 10, 'bold'))
        self.upper_label.grid(row=0, column=0, columnspan=2)

        self.button1 = Button(self.fr, text="Видалити обраного\n працівника", width=20, height=2, font=("arial", 10),
                              command=self.delete_employee)
        self.button1.grid(row=1, column=0)

        self.button2 = Button(self.fr, text="Конфігурація обраного\n працівника", width=20, height=2, font=("arial", 10),
                              command=self.sub_sub_configure_employee_menu)
        self.button2.grid(row=1, column=1)

    def sub_division_menu(self, event):
        pass

    def sub_sub_configure_employee_menu(self):
        selected_item = self.table.selection()[0]
        listed_raw = [str(i) for i in list(self.table.item(selected_item)['values'])]
        """
        name = listed_raw[1]
        surname = listed_raw[2]
        last_name = listed_raw[3]
        position = listed_raw[4]
        subdivision = listed_raw[5]
        """

        self.entry1 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry1.insert(0, listed_raw[1])
        self.entry1.grid(row=2, column=0)

        self.entry2 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry2.insert(0, listed_raw[2])
        self.entry2.grid(row=3, column=0)

        self.entry3 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry3.insert(0, listed_raw[3])
        self.entry3.grid(row=4, column=0)

        self.accept_but = Button(self.fr, text="Прийняти зміни", width=15, height=1, font=("arial", 10),
                              command=self.configure_employee)
        self.accept_but.grid(row=4, column=1)

        self.entry4 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry4.insert(0, listed_raw[4])
        self.entry4.grid(row=5, column=0)

        self.entry5 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry5.insert(0, listed_raw[5])
        self.entry5.grid(row=6, column=0)

    def delete_employee(self):
        selected_item = self.table.selection()[0]
        values = tuple(self.table.item(selected_item)['values'])
        self.db_access.del_employee(values)
        self.orders_generator.generate_firing_order(list(values)[1:])
        self.rt.destroy()
        self.table.delete(selected_item)

    def configure_employee(self):
        selected_item = self.table.selection()[0]
        listed_raw = [str(i) for i in list(self.table.item(selected_item)['values'])]
        values = tuple(self.table.item(selected_item)['values'])

        name = self.entry1.get()
        surname = self.entry2.get()
        last_name = self.entry3.get()
        position = self.entry4.get()
        subdivision = self.entry5.get()

        if listed_raw[1] == name and listed_raw[2] == surname and listed_raw[3] == last_name:
            if listed_raw[4] != position or listed_raw[5] != subdivision:
                self.orders_generator.generate_move_order(listed_raw[1:], [position, subdivision])

        self.db_access.del_employee(values)
        self.db_access.add_employee(name, surname, last_name, position, subdivision)
        self.table.item(selected_item, values=(listed_raw[0], name, surname, last_name, position, subdivision))
        self.rt.destroy()


class PyAccountingDataBase:
    def __init__(self, db_path='db.db'):
        self.db_path = db_path
        self.orders_generator = WordOrdedGenerator()

    def create_db(self):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute('''CREATE TABLE IF NOT EXISTS employees 
                     (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                      name TEXT NOT NULL, 
                      surname TEXT NOT NULL, 
                      last_name TEXT NOT NULL, 
                      position TEXT NOT NULL, 
                      subdivision TEXT NOT NULL)''')

        curs.execute('''CREATE TABLE IF NOT EXISTS subdivision 
                     (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                      name TEXT NOT NULL, 
                      e_number INTEGER NOT NULL)''')
        conn.commit()
        conn.close()

    def del_employee(self, values: tuple):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        query = "DELETE FROM employees WHERE id=? AND name=? AND surname=? AND last_name=? AND position=? AND subdivision=?"
        curs.execute(query, values)
        conn.commit()
        curs.execute("SELECT * FROM subdivision WHERE name=?", (values[5],))
        result = curs.fetchone()
        if result:
            conn.commit()
            curs.execute('UPDATE subdivision SET e_number=? WHERE name=?', (result[2] - 1, result[1]))
            conn.commit()
            curs.execute('SELECT * FROM subdivision WHERE name=?', (result[1],))
            result = curs.fetchone()
            conn.commit()
            if result[2] == 0:
                curs.execute('DELETE FROM subdivision WHERE e_number=?', (0,))
                conn.commit()
            else:
                pass
        conn.close()

    def add_employee(self, name, surname, last_name, position, subdivision):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute("SELECT * FROM subdivision WHERE name=?", (subdivision,))  # находим в БД отряд, в который хочет
        result = curs.fetchone()                                                # влезть наш новый работник
        conn.commit()

        if result:          # если такой отряд существует
            curs.execute("INSERT INTO employees (name, surname, last_name, position, subdivision) VALUES (?, ?, ?, ?, ?)",
                         (name, surname, last_name, position, subdivision))
            curs.execute("UPDATE subdivision SET e_number=? WHERE name=?", (result[2] + 1, subdivision))
            conn.commit()

        else:               # если такого отряда не существует
            curs.execute("INSERT INTO subdivision (name, e_number) VALUES (?, ?)", (subdivision, 1))
            curs.execute("INSERT INTO employees (name, surname, last_name, position, subdivision) VALUES (?, ?, ?, ?, ?)",
                         (name, surname, last_name, position, subdivision))
            conn.commit()

        conn.close()

    def add_subdivision(self, name):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute("INSERT INTO subdivision (name, e_number) VALUES (?, ?)", (name, 0))
        conn.commit()
        conn.close()

    def show_in_console(self):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute('SELECT * FROM employees')
        result = curs.fetchall()
        conn.commit()
        print('Employees table:')
        for raw in result:
            print(raw)
        print('---'*50)
        curs.execute('SELECT * FROM subdivision')
        result = curs.fetchall()
        conn.commit()
        print('Subdivisions table:')
        for raw in result:
            print(raw)
        print('---'*50)

    def show_employees(self):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute("SELECT * FROM employees")
        rows = tuple(curs.fetchall())
        conn.commit()
        root = Tk()
        root.geometry('{}x{}'.format(root.winfo_screenwidth(), root.winfo_screenheight()-100))
        root.title("Довідка з БД про працівників компанії")
        table = Table(root, headings=('ID', 'Ім\'я', 'Прізвище', 'По батькові', 'Посада', 'Відділ'), rows=rows)
        table.pack(expand=YES, fill=BOTH)
        root.mainloop()

    def show_subdivisions(self):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute("SELECT * FROM subdivision")
        rows = tuple(curs.fetchall())
        conn.commit()
        root = Tk()
        root.geometry('900x600')
        root.title("Довідка з БД про відділення компанії")
        table = Table(root, headings=('ID', 'Назва', 'Кількість робітників'), rows=rows)
        table.pack(expand=YES, fill=BOTH)
        root.mainloop()

    def add_employee_finally(self):
        name = self.name_entry.get()
        surname = self.surname_entry.get()
        last_name = self.last_name_entry.get()
        position = self.position_entry.get()
        subdivision = self.subdivision_entry.get()
        self.add_employee(name, surname, last_name, position, subdivision)
        self.orders_generator.generate_hiring_order([name, surname, last_name, position, subdivision])
        self.lil_root.destroy()

    def add_employee_menu(self):
        self.lil_root = Toplevel()
        self.lil_root.title('Додавання працівника')
        lil_frame = Frame(self.lil_root, bd=5)
        lil_frame.pack(fill=BOTH, expand=True, side=TOP)

        Label(lil_frame, text='Ім\'я: ', font=('arial', 10)).grid(row=2, column=0)
        self.name_entry = Entry(lil_frame, font=('arial', 10), bd=5)
        self.name_entry.grid(row=2, column=1)

        Label(lil_frame, text='Прізвище: ', font=('arial', 10)).grid(row=3, column=0)
        self.surname_entry = Entry(lil_frame, font=('arial', 10), bd=5)
        self.surname_entry.grid(row=3, column=1)

        Label(lil_frame, text='По батькові: ', font=('arial', 10)).grid(row=4, column=0)
        self.last_name_entry = Entry(lil_frame, font=('arial', 10), bd=5)
        self.last_name_entry.grid(row=4, column=1)

        Label(lil_frame, text='Посада: ', font=('arial', 10)).grid(row=5, column=0)
        self.position_entry = Entry(lil_frame, font=('arial', 10), bd=5)
        self.position_entry.grid(row=5, column=1)

        Label(lil_frame, text='Відділ: ', font=('arial', 10)).grid(row=6, column=0)
        self.subdivision_entry = Entry(lil_frame, font=('arial', 10), bd=5)
        self.subdivision_entry.grid(row=6, column=1)

        accept_but = Button(lil_frame, text="Додати працівника", width=15, height=1, font=("arial", 10, 'bold'),
                            command=self.add_employee_finally)
        accept_but.grid(row=4, column=2)

        self.lil_root.mainloop()


class WordOrdedGenerator:
    def __init__(self, hiring_pattern_path=r"PATTERNS\hiring_pattern.docx",
                 firing_pattern_path=r"PATTERNS\firing_pattern.docx",
                 move_pattern_path=r"PATTERNS\move_pattern.docx"):
        self.hiring_pattern_path = hiring_pattern_path
        self.firing_pattern_path = firing_pattern_path
        self.move_pattern_path = move_pattern_path

    def generate_firing_order(self, lst):
        """ lst = [name, surname, last_name, position, subdivision]"""
        #print('lst = ', lst)
        check_list = ['%NAME%', '%SURNAME%', '%LASTNAME%', '%POSITION%', '%SUBDIVISION%']
        document = Document(self.firing_pattern_path)

        for j in range(len(check_list)):
            ON_AIR, ON_GROUND = check_list[j], lst[j]

            for paragraph in document.paragraphs:
                if ON_AIR in paragraph.text:

                    chunks = paragraph.runs
                    for i in range(len(chunks)):
                        if ON_AIR in chunks[i].text:
                            text = chunks[i].text.replace(ON_AIR, ON_GROUND )
                            chunks[i].text = text

        pre_path = "ORDERS\{}. Наказ про звільнення {}-а {} {} {} з \'{}\'.docx".format(
            datetime.strftime(datetime.now(), '%H-%M-%S %d-%m-%Y'), lst[3], lst[1], lst[0], lst[2], lst[4])

        document.save(pre_path)

    def generate_hiring_order(self, lst):
        """ lst = [name, surname, last_name, position, subdivision]"""
        check_list = ['%NAME%', '%SURNAME%', '%LASTNAME%', '%POSITION%', '%SUBDIVISION%']
        document = Document(self.hiring_pattern_path)

        for j in range(len(check_list)):
            ON_AIR, ON_GROUND = check_list[j], lst[j]

            for paragraph in document.paragraphs:
                if ON_AIR in paragraph.text:

                    chunks = paragraph.runs
                    for i in range(len(chunks)):
                        if ON_AIR in chunks[i].text:
                            text = chunks[i].text.replace(ON_AIR, ON_GROUND )
                            chunks[i].text = text

        pre_path = "ORDERS\{}. Наказ про прийняття на позицію {}-а {} {} {} у відділ {}.docx".format(
        datetime.strftime(datetime.now(), '%H-%M-%S %d-%m-%Y'), lst[3], lst[1],lst[0], lst[2],lst[4])

        document.save(pre_path)

    def generate_move_order(self, old_lst, move_info_lst):
        """
            old_lst = [name, surname, last_name, old_position, old_subdivision]
            move_info_lst = [new_position, new_subdivision]
        """
        work_list = old_lst + move_info_lst

        check_list = ['%NAME%', '%SURNAME%', '%LASTNAME%', '%POSITION%', '%SUBDIVISION%', '%NEWPOS%', '%NEWSUB%']
        document = Document(self.move_pattern_path)

        for j in range(len(check_list)):
            ON_AIR, ON_GROUND = check_list[j], work_list[j]

            #print(ON_AIR, ON_GROUND)

            for paragraph in document.paragraphs:
                if ON_AIR in paragraph.text:

                    chunks = paragraph.runs
                    for i in range(len(chunks)):
                        if ON_AIR in chunks[i].text:
                            text = chunks[i].text.replace(ON_AIR, ON_GROUND)
                            chunks[i].text = text

        pre_path = "ORDERS\{}. Наказ про переведення працівника {} {} {}.docx".format(
        datetime.strftime(datetime.now(), '%H-%M-%S %d-%m-%Y'), work_list[1], work_list[0], work_list[2])

        document.save(pre_path)


class PyAccounting:
    def __init__(self, orders_path='ORDERS'):
        self.PAdb = PyAccountingDataBase()
        self.orders_path = orders_path

        self.root = Tk()
        self.root.title('PyAccounting')

        self.frame = Frame(self.root, bd=5)
        self.frame.pack(fill=BOTH, expand=True, side=TOP)

        self.button1 = Button(self.frame, text="Конфігурія обліку робітників", width=30, height=2,
                              font=("arial", 13, "bold"), command=self.PAdb.show_employees)
        self.button1.grid(row=0, column=0)

        self.button2 = Button(self.frame, text="Конфігурія підрозділів робітників", width=30, height=2,
                              font=("arial", 13, "bold"), command=self.PAdb.show_subdivisions)
        self.button2.grid(row=0, column=1)

        self.button3 = Button(self.frame, text="Додати нового працівника", width=30, height=2,
                              font=("arial", 13, "bold"), command=self.PAdb.add_employee_menu)
        self.button3.grid(row=1, column=0)

        self.button4 = Button(self.frame, text="Перегляд наказів", width=30, height=2, font=("arial", 13, "bold"),
                              command=self.open_dir_with_orders)
        self.button4.grid(row=1, column=1)

        self.root.mainloop()

    def open_dir_with_orders(self):
        os.startfile(self.orders_path)




