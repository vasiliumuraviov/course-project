import sqlite3
from tkinter import *
import tkinter.ttk as ttk
from docx import Document
from typing import List
import os
from datetime import datetime


class Table(Frame):
    def __init__(self, parent=None, headings=tuple(), rows=tuple()) -> None:
        super().__init__(parent)

        # личный интерфейс для работы с базой данных
        self.db_access = PyAccountingDataBase()
        # личный генератор отчётности
        self.orders_generator = WordOrdersGenerator()

        # общая структура таблицы всех работников
        self.table = ttk.Treeview(self, show="headings", selectmode="browse")
        self.table["columns"] = headings
        self.table["displaycolumns"] = headings

        for head in headings:
            self.table.heading(head, text=head, anchor=CENTER)
            self.table.column(head, anchor=CENTER)

        for row in rows:
            self.table.insert('', END, values=tuple(row))

        # биндим кнопки выпадающих подменю
        if headings[1] == 'Ім\'я':
            self.table.bind('<Double-Button-1>', self.sub_employee_menu)
        elif headings[1] == 'Назва':
            # кнопка-пустышка, так как никакого дополнительного
            # функционала для меню отделов комании нету в требованиях к проекту
            self.table.bind('<Double-Button-1>', self.sub_division_menu)

        # последние конфигурации структуры таблицы всех работников
        scrolltable = Scrollbar(self, command=self.table.yview)
        self.table.configure(yscrollcommand=scrolltable.set)
        scrolltable.pack(side=RIGHT, fill=Y)
        self.table.pack(expand=YES, fill=BOTH)

    # метод реализовывающий подменю редактирования/удаления информации о работниках
    def sub_employee_menu(self, event):
        self.rt = Toplevel(self.table)
        self.rt.title('Конфігурація працівника')

        # выбраная пользователем строка(работник) в таблице сотрудников
        selected_item = self.table.selection()[0]

        # форматируем выбранные данные для удобного отображения
        listed_raw = [str(i) for i in list(self.table.item(selected_item)['values'])]
        full_name, position, subdivision = ' '.join(listed_raw[1:4]), listed_raw[4], listed_raw[5]
        res = f"- {full_name};\n- {position};\n- {subdivision};"

        # реализация меню типа " Что пожелаете сделать с сотрудником 'Васильев Василий Васильевич'? "
        self.fr = Frame(self.rt, bd=5)
        self.fr.pack(fill=BOTH, expand=True, side=TOP)

        self.upper_label = Label(self.fr,
                                 text=f"Меню робітника:\n{res}",
                                 font=('arial', 10, 'bold'))
        self.upper_label.grid(row=0, column=0, columnspan=2)

        self.button1 = Button(self.fr,
                              text="Видалити обраного\n працівника",
                              width=20,
                              height=2,
                              font=("arial", 10),
                              command=self.delete_employee)
        self.button1.grid(row=1, column=0)

        self.button2 = Button(self.fr,
                              text="Конфігурація обраного\n працівника",
                              width=20,
                              height=2,
                              font=("arial", 10),
                              command=self.sub_sub_configure_employee_menu)
        self.button2.grid(row=1, column=1)

    # выше упомянутый в комментариях метод пустышка
    # (при желании можно добавить свою реализацию)
    def sub_division_menu(self, event):
        pass

    # подменю выпадающие после выбора пункта в меню по типу "конфигурация инфо. о сотруднике"
    def sub_sub_configure_employee_menu(self):
        # выбраная пользователем строка(работник) в таблице сотрудников
        selected_item = self.table.selection()[0]

        # вытягиваем по частям информацию о работнике из выбраной строки меню
        listed_raw = [str(i) for i in list(self.table.item(selected_item)['values'])]

        # поле редактирования имени
        self.entry1 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry1.insert(0, listed_raw[1])
        self.entry1.grid(row=2, column=0)

        # поле редактирования фамилии
        self.entry2 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry2.insert(0, listed_raw[2])
        self.entry2.grid(row=3, column=0)

        # поле редактирования отчества
        self.entry3 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry3.insert(0, listed_raw[3])
        self.entry3.grid(row=4, column=0)

        # поле редактирования должности
        self.entry4 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry4.insert(0, listed_raw[4])
        self.entry4.grid(row=5, column=0)

        # поле редактирования отдела
        self.entry5 = Entry(self.fr, font=('arial', 10), bd=5)
        self.entry5.insert(0, listed_raw[5])
        self.entry5.grid(row=6, column=0)

        # кнопка подтверждения
        self.accept_but = Button(self.fr, text="Прийняти зміни", width=15, height=1, font=("arial", 10),
                                 command=self.configure_employee)
        self.accept_but.grid(row=4, column=1)

    # метод реализующий меню удаления сотрудника
    def delete_employee(self):
        # выбраная пользователем строка(работник) в таблице сотрудников
        selected_item = self.table.selection()[0]

        # формируем кортеж для использования как параметр в методе удаления сотрудника из БД
        values = tuple(self.table.item(selected_item)['values'])

        # удаляем сотрудника из БД
        self.db_access.del_employee(values)

        # генерируем отчёт об увольнении в папку ORDERS
        self.orders_generator.generate_firing_order(list(values)[1:])

        # закрываем подменю редактирования сотрудника(так как мы его удалили)
        self.rt.destroy()

        # удаляем строку с инфо. о сотруднике из общей таблицы в меню
        self.table.delete(selected_item)

    # метод реализующий подтверждение изменений информации о сотруднике
    def configure_employee(self):
        # выбраная пользователем строка(работник) в таблице сотрудников
        selected_item = self.table.selection()[0]

        # вытягиваем по частям информацию о работнике из выбраной строки меню
        listed_raw = [str(i) for i in list(self.table.item(selected_item)['values'])]

        # формируем кортеж для использования как параметр в методе удаления сотрудника из БД
        values = tuple(self.table.item(selected_item)['values'])

        # вытягиваем данные из заполненных пользователем полей с новой инфо. о сотруднике
        new_name = self.entry1.get()
        new_surname = self.entry2.get()
        new_last_name = self.entry3.get()
        new_position = self.entry4.get()
        new_subdivision = self.entry5.get()

        # проверка неизменности ФИО, так как такие изменения не документируются, а
        # считаются за простые исправления ошибок правильности введённых данных
        if listed_raw[1] == new_name and \
            listed_raw[2] == new_surname and \
                listed_raw[3] == new_last_name:

            # если ФИО не изменилось, значит работника перевели в новый отдел и/или повысили/понизили
            if listed_raw[4] != new_position or listed_raw[5] != new_subdivision:
                # генерацию соответствующего приказа о повышении и/или переводе
                self.orders_generator.generate_move_order(listed_raw[1:], [new_position, new_subdivision])


        # удаляем сотрудника из БД
        self.db_access.del_employee(values)

        # добавляем нового с изменёнными параметрами
        self.db_access.add_employee(new_name, new_surname, new_last_name, new_position, new_subdivision)

        # добавляем нового сотрудника в отображаемую таблицу
        self.table.item(selected_item,
                        values=(listed_raw[0], new_name, new_surname, new_last_name, new_position, new_subdivision))

        # закрываем подменю
        self.rt.destroy()


# простой класс реализующий методы для работы с БД сотрудников
# посредством выполнения простого SQL кода.
class PyAccountingDataBase:
    def __init__(self, db_path='db.db'):
        self.db_path = db_path
        self.orders_generator = WordOrdersGenerator()

    # метод создающий БД
    def create_db(self):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute('''CREATE TABLE IF NOT EXISTS employees 
                     (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                      name TEXT NOT NULL, 
                      surname TEXT NOT NULL, 
                      last_name TEXT NOT NULL, 
                      position TEXT NOT NULL, 
                      subdivision TEXT NOT NULL)''')

        curs.execute('''CREATE TABLE IF NOT EXISTS subdivision 
                     (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                      name TEXT NOT NULL, 
                      e_number INTEGER NOT NULL)''')
        conn.commit()
        conn.close()

    # метод удаляющий сотрудника из БД
    def del_employee(self, values: tuple):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()

        query = "DELETE FROM employees WHERE id=? AND name=? AND surname=? AND last_name=? AND position=? AND subdivision=?"
        curs.execute(query, values)
        conn.commit()

        curs.execute("SELECT * FROM subdivision WHERE name=?", (values[5],))
        result = curs.fetchone()

        if result:
            conn.commit()

            curs.execute('UPDATE subdivision SET e_number=? WHERE name=?', (result[2] - 1, result[1]))
            conn.commit()

            curs.execute('SELECT * FROM subdivision WHERE name=?', (result[1],))
            result = curs.fetchone()
            conn.commit()

            if result[2] == 0:
                curs.execute('DELETE FROM subdivision WHERE e_number=?', (0,))
                conn.commit()
            else:
                pass

        conn.close()

    # метод добавляющий сотрудника в БД
    def add_employee(self, name: str, surname: str, last_name: str, position: str, subdivision: str) -> None:
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute("SELECT * FROM subdivision WHERE name=?", (subdivision,))  # находим в БД отряд, в который хочет
        result = curs.fetchone()                                                # влезть наш новый работник
        conn.commit()

        if result:          # если такой отряд существует
            curs.execute("INSERT INTO employees (name, surname, last_name, position, subdivision) VALUES (?, ?, ?, ?, ?)",
                         (name, surname, last_name, position, subdivision))
            curs.execute("UPDATE subdivision SET e_number=? WHERE name=?", (result[2] + 1, subdivision))
            conn.commit()

        else:               # если такого отряда не существует
            curs.execute("INSERT INTO subdivision (name, e_number) VALUES (?, ?)", (subdivision, 1))
            curs.execute("INSERT INTO employees (name, surname, last_name, position, subdivision) VALUES (?, ?, ?, ?, ?)",
                         (name, surname, last_name, position, subdivision))
            conn.commit()

        conn.close()

    # метод добавляющий отдел в БД
    def add_subdivision(self, name):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute("INSERT INTO subdivision (name, e_number) VALUES (?, ?)", (name, 0))
        conn.commit()
        conn.close()

    # метод для отображения всех сотрудников в консоли
    def _show_in_console(self) -> None:
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute('SELECT * FROM employees')
        result = curs.fetchall()
        conn.commit()
        print('Employees table:')
        for raw in result:
            print(raw)
        print('---'*50)
        curs.execute('SELECT * FROM subdivision')
        result = curs.fetchall()
        conn.commit()
        print('Subdivisions table:')
        for raw in result:
            print(raw)
        print('---'*50)

    # метод показа таблицы всех сотрудников с их информацией
    def show_employees(self):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute("SELECT * FROM employees")
        rows = tuple(curs.fetchall())
        conn.commit()
        root = Tk()
        root.geometry('{}x{}'.format(root.winfo_screenwidth(), root.winfo_screenheight()-100))
        root.title("Довідка з БД про працівників компанії")
        table = Table(root, headings=('ID', 'Ім\'я', 'Прізвище', 'По батькові', 'Посада', 'Відділ'), rows=rows)
        table.pack(expand=YES, fill=BOTH)
        root.mainloop()

    # метод показа таблицы всех подразделений с их информацией
    def show_subdivisions(self):
        conn = sqlite3.connect(self.db_path)
        curs = conn.cursor()
        curs.execute("SELECT * FROM subdivision")
        rows = tuple(curs.fetchall())
        conn.commit()
        root = Tk()
        root.geometry('900x600')
        root.title("Довідка з БД про відділення компанії")
        table = Table(root, headings=('ID', 'Назва', 'Кількість робітників'), rows=rows)
        table.pack(expand=YES, fill=BOTH)
        root.mainloop()

    # метод, привязанный к кнопке подтверждения "найма" сотрудника. полностью обновляет БД и меню
    def add_employee_finally(self):
        name = self.first_name_entry.get()
        surname = self.surname_entry.get()
        last_name = self.last_name_entry.get()
        position = self.position_entry.get()
        subdivision = self.subdivision_entry.get()

        self.add_employee(name, surname, last_name, position, subdivision)
        self.orders_generator.generate_hiring_order([name, surname, last_name, position, subdivision])

        self.hire_menu_root.destroy()

    # метод создаёт меню "найма"(по факту, создания) нового работника
    def add_employee_menu(self) -> None:
        self.hire_menu_root = Toplevel()
        self.hire_menu_root.title('Додавання працівника')

        hire_menu_frame = Frame(self.hire_menu_root, bd=5)
        hire_menu_frame.pack(fill=BOTH, expand=True, side=TOP)

        Label(hire_menu_frame,
              text='Ім\'я: ',
              font=('arial', 10)).grid(row=2, column=0)
        self.first_name_entry = Entry(hire_menu_frame, font=('arial', 10), bd=5)
        self.first_name_entry.grid(row=2, column=1)

        Label(hire_menu_frame,
              text='Прізвище: ',
              font=('arial', 10)).grid(row=3, column=0)
        self.surname_entry = Entry(hire_menu_frame, font=('arial', 10), bd=5)
        self.surname_entry.grid(row=3, column=1)

        Label(hire_menu_frame,
              text='По батькові: ',
              font=('arial', 10)).grid(row=4, column=0)
        self.last_name_entry = Entry(hire_menu_frame, font=('arial', 10), bd=5)
        self.last_name_entry.grid(row=4, column=1)

        Label(hire_menu_frame,
              text='Посада: ',
              font=('arial', 10)).grid(row=5, column=0)
        self.position_entry = Entry(hire_menu_frame, font=('arial', 10), bd=5)
        self.position_entry.grid(row=5, column=1)

        Label(hire_menu_frame,
              text='Відділ: ',
              font=('arial', 10)).grid(row=6, column=0)
        self.subdivision_entry = Entry(hire_menu_frame, font=('arial', 10), bd=5)
        self.subdivision_entry.grid(row=6, column=1)

        accept_but = Button(hire_menu_frame,
                            text="Додати працівника",
                            width=15,
                            height=1,
                            font=("arial", 10, 'bold'),
                            command=self.add_employee_finally)
        accept_but.grid(row=4, column=2)

        self.hire_menu_root.mainloop()


class WordOrdersGenerator:
    # класс, который помогает генерировать .docx указы
    def __init__(self, hiring_pattern_path=r"PATTERNS\hiring_pattern.docx",
                 firing_pattern_path=r"PATTERNS\firing_pattern.docx",
                 move_pattern_path=r"PATTERNS\move_pattern.docx") -> None:

        self.hiring_pattern_path: str = hiring_pattern_path
        self.firing_pattern_path: str = firing_pattern_path
        self.move_pattern_path: str = move_pattern_path

    # метод генерирующий .docx указ об увольнении сотрудника
    def generate_firing_order(self, lst: List[str]) -> None:
        """ lst = [name, surname, last_name, position, subdivision]"""

        check_list: List[str] = ['%NAME%', '%SURNAME%', '%LASTNAME%', '%POSITION%', '%SUBDIVISION%']
        document: Document = Document(self.firing_pattern_path)

        for j in range(len(check_list)):
            curr_repl, curr_sub = check_list[j], lst[j]

            for paragraph in document.paragraphs:
                if curr_repl in paragraph.text:

                    chunks = paragraph.runs
                    for i in range(len(chunks)):
                        if curr_repl in chunks[i].text:
                            text = chunks[i].text.replace(curr_repl, curr_sub )
                            chunks[i].text = text

        current_time: str = datetime.strftime(datetime.now(), '%H-%M-%S %d-%m-%Y')
        pre_path = f"ORDERS\\{current_time}. " \
            f"Наказ про звільнення {lst[3]}-а {lst[1]} {lst[0]} {lst[2]} з \'{lst[4]}\'.docx"

        document.save(pre_path)

    # метод генерирующий .docx указ о найме сотрудника
    def generate_hiring_order(self, lst: List[str]) -> None:
        """ lst = [name, surname, last_name, position, subdivision]"""

        check_list: List[str] = ['%NAME%', '%SURNAME%', '%LASTNAME%', '%POSITION%', '%SUBDIVISION%']

        document: Document = Document(self.hiring_pattern_path)

        for j in range(len(check_list)):
            curr_repl, curr_sub = check_list[j], lst[j]

            for paragraph in document.paragraphs:
                if curr_repl in paragraph.text:

                    chunks = paragraph.runs
                    for i in range(len(chunks)):
                        if curr_repl in chunks[i].text:
                            text = chunks[i].text.replace(curr_repl, curr_sub)
                            chunks[i].text = text

        curr_time: str = datetime.strftime(datetime.now(), '%H-%M-%S %d-%m-%Y')
        pre_path: str = f"ORDERS\\{curr_time}. " \
            f"Наказ про прийняття на позицію {lst[3]}-а {lst[1]} {lst[0]} {lst[2]} у відділ {lst[4]}.docx"

        document.save(pre_path)

    # метод генерирующий .docx указ о переводе и/или повышении/понижении сотрудника
    def generate_move_order(self, curr_employee_info: List[str], move_info_from_to: List[str]) -> None:
        """
            old_lst = [name, surname, last_name, old_position, old_subdivision]
            move_info_lst = [new_position, new_subdivision]
        """
        sub_info_lst: list = curr_employee_info + move_info_from_to
        check_list: list = ['%NAME%', '%SURNAME%', '%LASTNAME%', '%POSITION%', '%SUBDIVISION%', '%NEWPOS%', '%NEWSUB%']

        document = Document(self.move_pattern_path)

        for j in range(len(check_list)):
            curr_repl, curr_sub = check_list[j], sub_info_lst[j]

            for paragraph in document.paragraphs:
                if curr_repl in paragraph.text:

                    chunks = paragraph.runs
                    for i in range(len(chunks)):
                        if curr_repl in chunks[i].text:
                            text = chunks[i].text.replace(curr_repl, curr_sub)
                            chunks[i].text = text

        current_time_str: str = datetime.strftime(datetime.now(), '%H-%M-%S %d-%m-%Y')
        pre_path: str = f"ORDERS\\{current_time_str}." \
            f" Наказ про переведення працівника {sub_info_lst[1]} {sub_info_lst[0]} {sub_info_lst[2]}.docx"

        document.save(pre_path)


class PyAccounting:
    def __init__(self, orders_path='ORDERS') -> None:
        # конструктор инциализирует главное меню, кнопки в нём и всё нужное для их корректной работы

        self.PAdb = PyAccountingDataBase()
        self.orders_path = orders_path

        self.root = Tk()
        self.root.title('PyAccounting')

        self.frame = Frame(self.root, bd=5)
        self.frame.pack(fill=BOTH, expand=True, side=TOP)

        self.button1 = Button(self.frame,
                              text="Конфігурія обліку робітників",
                              width=30,
                              height=2,
                              font=("arial", 13, "bold"),
                              command=self.PAdb.show_employees)
        self.button1.grid(row=0, column=0)

        self.button2 = Button(self.frame,
                              text="Конфігурія підрозділів робітників",
                              width=30,
                              height=2,
                              font=("arial", 13, "bold"),
                              command=self.PAdb.show_subdivisions)
        self.button2.grid(row=0, column=1)

        self.button3 = Button(self.frame,
                              text="Додати нового працівника",
                              width=30,
                              height=2,
                              font=("arial", 13, "bold"),
                              command=self.PAdb.add_employee_menu)
        self.button3.grid(row=1, column=0)

        self.button4 = Button(self.frame,
                              text="Перегляд наказів",
                              width=30,
                              height=2,
                              font=("arial", 13, "bold"),
                              command=self.open_dir_with_orders)
        self.button4.grid(row=1, column=1)

    # простой метод - открывает папку ORDERS с приказами в формате .docx
    def open_dir_with_orders(self) -> None:
        os.startfile(self.orders_path)

    def start(self) -> None:
        self.root.mainloop()


if __name__ == '__main__':
    session = PyAccounting()
    session.start()
